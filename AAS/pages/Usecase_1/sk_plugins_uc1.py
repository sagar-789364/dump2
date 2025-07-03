# Dependencies and imports
import os
import asyncio
import uuid
import json
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from dotenv import load_dotenv

# Semantic Kernel imports
from semantic_kernel import Kernel
from semantic_kernel.connectors.ai.open_ai import AzureTextEmbedding
from semantic_kernel.functions import kernel_function
from semantic_kernel.connectors.ai.function_choice_behavior import FunctionChoiceBehavior
from openai import AzureOpenAI
from azure.search.documents import IndexDocumentsBatch
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    SearchIndex,
    SimpleField,
    SearchableField,
    SearchFieldDataType,
    VectorSearch,
    HnswAlgorithmConfiguration,
    VectorSearchProfile,
    SearchField,
    SemanticField,
    SemanticSearch,
    SemanticPrioritizedFields,
    SemanticConfiguration
)
from semantic_kernel.contents.chat_history import ChatHistory
from semantic_kernel.connectors.ai.open_ai import AzureChatCompletion
from semantic_kernel.functions.kernel_arguments import KernelArguments
from semantic_kernel.connectors.ai.open_ai.prompt_execution_settings.azure_chat_prompt_execution_settings import (
    AzureChatPromptExecutionSettings,
)

# Azure imports
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.models import VectorizedQuery
from docx import Document
from io import BytesIO

# Load the .env file
load_dotenv()

# Configuration and data models
@dataclass
class AzureConfig:
    search_endpoint = os.getenv("search_endpoint")
    search_api_key = os.getenv("search_api_key")
    file_index_name = os.getenv("file_index_name")  # Use env var for chat/session memory
    memory_index_name = os.getenv("memory_index_name")  # Use env var for chat/session memory
    
    openai_api_version = os.getenv("openai_api_version")
    openai_api_key = os.getenv("openai_api_key")
    openai_endpoint = os.getenv("openai_endpoint")
        
    chat_deployment = "gpt-4o-mini"
    embedding_deployment = "text-embedding-3-large"

# Document Retrieval Plugin - Handle Both String and JSON
class DocumentRetrievalPlugin:
    """Plugin that handles both string and JSON array inputs, now using 'author' fields and correct filter order."""
    
    def __init__(self, config: AzureConfig):
        self.config = config
        self.search_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.file_index_name,
            credential=AzureKeyCredential(config.search_api_key)
        )
        self.embedding_service = AzureTextEmbedding(
            deployment_name=config.embedding_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )

        
    async def get_embedding(self, text: str) -> List[float]:
        """Get embedding for text using Azure OpenAI"""
        try:
            result = await self.embedding_service.generate_embeddings([text])
            return result[0].data
        except Exception as e:
            print(f"Warning: Azure embedding generation failed: {e}")
            return [0.0] * 3072
        
    def calculate_cosine_similarity(self, embedding1: List[float], embedding2: List[float]) -> float:
        """Calculate cosine similarity between two embeddings"""
        import numpy as np
    
        # Convert to numpy arrays
        vec1 = np.array(embedding1)
        vec2 = np.array(embedding2)
    
        # Calculate cosine similarity
        cos_sim = np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))
        return float(cos_sim)  # Convert to float for JSON serialization
        
    @kernel_function(
        description="Retrieve relevant documents based on username, framework, accounting_standards, topics, author types, scenario, query, and new fields",
        name="retrieve_documents"
    )
    async def retrieve_documents(
        self,
        kernel: Kernel,
        arguments: KernelArguments
    ) -> str:
        """Fixed retrieve documents - handles both string and JSON array inputs"""
        try:
            # user = arguments.get("user", None)
            framework = arguments.get("framework", None)
            accounting_standards = arguments.get("accounting_standards", None)
            topics = arguments.get("topics", None)
            author = arguments["author"]
            scenario = arguments["scenario"]
            query = arguments["query"]
            top_k = arguments.get("top_k", 5)
            scenario_vector = await self.get_embedding(scenario)
            filter_clauses = []
            # if user:
            #     filter_clauses.append(f"user eq '{user}'")
            #     print("[DEBUG] File index filter: user applied")
            if framework:
                filter_clauses.append(f"framework eq '{framework}'")
                print("[DEBUG] File index filter: framework applied")
            if accounting_standards:
                filter_clauses.append(f"accounting_standards eq '{accounting_standards}'")
                print("[DEBUG] File index filter: accounting_standards applied")
            if topics:
                filter_clauses.append(f"topics eq '{topics}'")
                print("[DEBUG] File index filter: topics applied")
            filter_clauses.append(f"author eq '{author}'")
            print("[DEBUG] File index filter: author applied")
            print(f"[DEBUG] File index filter order: framework, accounting_standards, topics, author")
            filter_query = " and ".join(filter_clauses)

            vector_query = VectorizedQuery(
                vector=scenario_vector,
                kind="vector",
                fields="embedding",
                k_nearest_neighbors=top_k
            )

            results = self.search_client.search(
                search_text=query,
                filter=filter_query,
                vector_queries=[vector_query],
                top=top_k,
                query_type="semantic",
                semantic_configuration_name="guidance-files-semantic-config",
                query_language="en-us",
                search_fields=["text"],
                search_mode="all"
            )

            documents = []
            for result in results:
                # Extract the vector and semantic scores
                # vector_score = result.get('@search.score', 0)
                result_vector = await self.get_embedding(result.get("text", ""))
                vector_score = self.calculate_cosine_similarity(scenario_vector, result_vector)
                semantic_score = result.get('@search.reranker_score', 0)
                semantic_score /= 4  # Normalize to [0, 1]
                combined_score = (vector_score * 0.7) + (semantic_score * 0.3)
                
                print("Score of Document Retrieval from Original Source: ")
                print(f"Vector Score: {vector_score}")
                print(f"Semantic Score: {semantic_score}")
                print(f"Combined Score: {combined_score}")

                # Add threshold check
                if combined_score > 0.4:  # You can adjust this threshold
                    documents.append({
                        "text": result.get("text", ""),
                        "source_file": result.get("source_file", "Unknown"),
                        "page": result.get("page", "N/A"),
                        "author": result.get("author", "Unknown"),
                        "relevance_score": combined_score
                    })
            # Check if any relevant documents were found
            if not documents:
                return json.dumps({"no_relevant_info": True})
            
            return json.dumps(documents)
            
        except Exception as e:
            error_msg = f"Retrieval error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})



# Response Generation Plugin
class ResponseGenerationPlugin:
    @kernel_function(
        description="Generate structured accounting guidance response",
        name="generate_response"
    )
    async def generate_response(
        self,
        kernel: Kernel,
        arguments: KernelArguments
    ) -> str:
        try:
            documents_json = arguments["documents_json"]
            scenario = arguments["scenario"]
            query = arguments["query"]
            framework = arguments.get("framework", "")
            accounting_standards = arguments.get("accounting_standards", "")
            topics = arguments.get("topics", "")
            
            documents = json.loads(documents_json)
            # Check for no relevant information
            if isinstance(documents, dict):
                if documents.get("no_relevant_info"):
                    return "Unable to Answer: No relevant information found in the knowledge base for this query."
                if "error" in documents:
                    return json.dumps({"error": documents["error"]})
            
            # If no documents found
            if not documents or (isinstance(documents, list) and len(documents) == 0):
                return "Unable to Answer: No relevant information found in the knowledge base for this query."
            
            doc_section = "\n".join([
                f"Source: {doc['source_file']}, Page: {doc['page']}\nText: {doc['text']}\n"
                for doc in documents
            ])
            
            history = ChatHistory()
            system_message = f"""
            You are a senior technical accounting expert, advisor and author. You are given a scenario, a query and the type of author(s). The type of author(s) may be single type or even multiple types. Based on the scenario and query given, you should analyze the scenario, the type(s) of authors, use the guidance files of those author(s) that are already given as indexes to you, provide a structured, detailed response. Your response must include:

*Framework*: {framework}
*Accounting Standards*: {accounting_standards}
*Topics*: {topics}

1. *Guidance*: Interpret the accounting treatment with references to authoritative guidance.
2. *Sample Filings*: Mention any relevant examples or patterns from filings.
3. *Sources*: The list in the document documents used, with page numbers and highlighted text. The text should be detailed as it is from the source.

Be precise, use professional accounting language, include all the relevant data in the response and ensure clarity for authors and controllers.

For each type of author you retrieve the data related to that type of authors only. For example: 
1. If the author is 'X', you would get the response related to the 'X' only not from the other type of author related files. The template of the response is as below:
        ## Guidance:  \n
        [Response you get from the indexes related to 'X' for guidance]
        ## Sample Filings:  \n
        [Response you get from the indexes related to 'X' for sample filings]
        ## Sources:  \n
        [Response you get from the indexes related to 'X' of sources]


3. If there are more than 2 types of different authors, then the responses related to all the different authors must be given in the similar pattern.
4. Don't consider any other authors other than those that are given. Don't even provide the heading with 'Response from 'X' author and so.

---
**Scenario**: {scenario}

**Question**: {query}

**Retrieved Documents**:
{doc_section}

-------------
**Response**:
            """
            history.add_system_message(system_message)
            history.add_user_message("Please provide the accounting guidance based on the given information.")
            
            # Create execution settings if not provided in arguments
            execution_settings = arguments.get("settings", AzureChatPromptExecutionSettings(
                temperature=0.0,
                max_tokens=2000
            ))
            
            # Get chat service and generate response
            chat_service = kernel.get_service(type=AzureChatCompletion)
            completion = await chat_service.get_chat_message_content(
                chat_history=history,
                settings=execution_settings,
                kernel=kernel
            )
            
            # Extract and return the actual response content
            if hasattr(completion, 'content'):
                return completion.content
            elif isinstance(completion, dict):
                return completion.get('content', '')
            else:
                return str(completion)
            
        except Exception as e:
            error_msg = f"Response generation failed: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg



# Document Export Plugin
class DocumentExportPlugin:
    """Plugin that formats Word documents and returns bytes for download"""
    
    @kernel_function(
        description="Save the generated response to a Word document with proper formatting and return bytes",
        name="save_to_docx"
    )
    def save_to_docx(
        self,
        kernel: Kernel,
        arguments: KernelArguments
        ) -> str:
        try:
            response = arguments["response"]
            scenario = arguments["scenario"]
            query = arguments["query"]
            filename = arguments["filename"]
            # Get new fields if present
            framework = arguments.get("framework", "")
            accounting_standards = arguments.get("accounting_standards", "")
            topics = arguments.get("topics", "")
            
            buffer = BytesIO()
            doc = Document()
            
            # Main title
            doc.add_heading("Accounting Guidance Report", level=1)
            
            # Add metadata section with new fields
            doc.add_heading("Metadata", level=2)
            doc.add_paragraph(f"Framework: {framework}")
            doc.add_paragraph(f"Accounting Standards: {accounting_standards}")
            doc.add_paragraph(f"Topics: {topics}")

            # Scenario section
            doc.add_heading("Scenario", level=2)
            doc.add_paragraph(scenario)
            
            # Query section  
            doc.add_heading("Query", level=2)
            doc.add_paragraph(query)
            
            # Response section
            doc.add_heading("Response", level=2)
            
            # Process the response text line by line to maintain formatting
            lines = response.split("\n")
            current_paragraph = ""
            
            for line in lines:
                line = line.strip()
                
                if not line:  # Empty line
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    continue
                
                # Handle different heading levels
                if line.startswith("#### "):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line[5:], level=4)
                elif line.startswith("### "):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line[2:], level=1)
                # Handle bold text markers
                elif line.startswith("*") and line.endswith("*") and len(line) > 2:
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    p = doc.add_paragraph()
                    p.add_run(line[1:-1]).bold = True
                # Handle bullet points
                elif line.startswith("- "):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_paragraph(line[2:], style='List Bullet')
                # Handle numbered lists
                elif any(line.startswith(f"{i}. ") for i in range(1, 10)):
                    if current_paragraph:
                        doc.add_paragraph(current_paragraph)
                        current_paragraph = ""
                    doc.add_paragraph(line[3:], style='List Number')
                # Regular text
                else:
                    if current_paragraph:
                        current_paragraph += " " + line
                    else:
                        current_paragraph = line
            
            # Add any remaining paragraph
            if current_paragraph:
                doc.add_paragraph(current_paragraph)
            
            doc.save(buffer)
            buffer.seek(0)
            
            import base64
            doc_bytes = buffer.read()
            doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')

            return json.dumps({
                "status": "success",
                "document_data": doc_base64,
                "filename": filename
            })
            
        except Exception as e:
            error_msg = f"Document generation failed: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})



class ConversationMemoryPlugin:
    """Combined plugin for conversation memory management and retrieval"""
    
    def __init__(self, config: AzureConfig):
        print("Initializing ConversationMemoryPlugin...")
        # Initialize Search clients
        self.index_client = SearchIndexClient(
            endpoint=config.search_endpoint,
            credential=AzureKeyCredential(config.search_api_key)
        )
        
        self.search_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.memory_index_name,
            credential=AzureKeyCredential(config.search_api_key)
        )

        self.openai_client = AzureOpenAI(
            api_key=config.openai_api_key,
            api_version=config.openai_api_version,
            azure_endpoint=config.openai_endpoint
        )
        
        # Ensure index exists
        self.ensure_index_exists()
    
    def ensure_index_exists(self, config=AzureConfig):
        """Create index if it doesn't exist"""
        try:
            print(f"Checking if index {config.memory_index_name} exists...")
            try:
                self.index_client.get_index(config.memory_index_name)
                print(f"Index {config.memory_index_name} exists")
                return  # Exit if index exists
            except:
                print(f"Index {config.memory_index_name} not found. Creating...")

            # Define semantic configuration
            semantic_config = SemanticConfiguration(
                name="my-semantic-config",
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="authors"),
                    keywords_fields=[
                        SemanticField(field_name="scenario"),
                        SemanticField(field_name="query"),
                    ],  
                    content_fields=[
                        SemanticField(field_name="response")
                    ]
                )
            )

            # Define vector search configuration
            vector_search = VectorSearch(
                algorithms=[
                    HnswAlgorithmConfiguration(
                        name="vector-config",
                        kind="hnsw",
                        parameters={
                            "m": 4,
                            "efConstruction": 400,
                            "efSearch": 500,
                            "metric": "cosine"
                        }
                    )
                ],
                profiles=[
                    VectorSearchProfile(
                        name="vector-profile",
                        algorithm_configuration_name="vector-config",
                    )
                ]
            )

            # Define fields
            fields = [
                SimpleField(name="id", type=SearchFieldDataType.String, key=True),
                SearchableField(name="user", type=SearchFieldDataType.String, filterable=True, searchable=True),
                SearchableField(name="framework", type=SearchFieldDataType.String, filterable=True, searchable=True),
                SearchableField(name="accounting_standards", type=SearchFieldDataType.String, filterable=True, searchable=True),
                SearchableField(name="topics", type=SearchFieldDataType.String, filterable=True, searchable=True),
                # Use a collection for authors (multi-author support)
                SearchableField(name="authors", type=SearchFieldDataType.String, filterable=True, searchable=True),
                SearchableField(name="scenario", type=SearchFieldDataType.String, searchable=True, analyzer_name="standard.lucene"),
                SearchableField(name="query", type=SearchFieldDataType.String, searchable=True, analyzer_name="standard.lucene"),
                SearchableField(name="response", type=SearchFieldDataType.String),
                SimpleField(name="timestamp", type=SearchFieldDataType.DateTimeOffset, filterable=True),
                SearchField(
                    name="embedding",
                    type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                    vector_search_dimensions=3072,
                    vector_search_profile_name="vector-profile"
                )
            ]

            # Create index
            index = SearchIndex(
                name=config.memory_index_name,
                fields=fields,
                vector_search=vector_search,
                semantic_search=SemanticSearch(configurations=[semantic_config])
            )

            print(f"Creating {config.memory_index_name}...")
            self.index_client.create_index(index)
            print("Index created successfully with semantic configuration")

        except Exception as e:
            print(f"Error ensuring index exists: {str(e)}")
            raise

    async def get_embedding(self, text: str, model = AzureConfig.embedding_deployment) -> list:
        response = self.openai_client.embeddings.create(
            input=[text],
            model=model
        )
        return response.data[0].embedding
    

    def calculate_cosine_similarity(self, embedding1: List[float], embedding2: List[float]) -> float:
        """Calculate cosine similarity between two embeddings"""
        import numpy as np
    
        # Convert to numpy arrays
        vec1 = np.array(embedding1)
        vec2 = np.array(embedding2)
    
        # Calculate cosine similarity
        cos_sim = np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))
        return float(cos_sim)  # Convert to float for JSON serialization

    @kernel_function(
    description="Export conversation memory to Azure Search",
    name="export_memory"
    )
    async def export_memory(
        self,
        kernel: Kernel,
        arguments: KernelArguments
        ) -> str:
        import json
        import uuid
        try:
            print("[DEBUG] Exporting conversation to memory index...")
            print(f"[DEBUG] Arguments: {arguments}")
            user = arguments.get("user", "")
            framework = arguments.get("framework", "")
            accounting_standards = arguments.get("accounting_standards", "")
            topics = arguments.get("topics", "")
            authors = arguments["authors"]
            scenario = arguments["scenario"]
            query = arguments["query"]
            response = arguments["response"]

            conversation_text = f"{scenario} {query}"
            embedding = await self.get_embedding(conversation_text)
            print(f"[DEBUG] Embedding generated (first 5 values)")

            # Always store authors as a comma-separated string (to match index schema)
            if isinstance(authors, str):
                try:
                    author_list = json.loads(authors)
                    if not isinstance(author_list, list):
                        author_list = [author_list]
                    authors_str = ','.join(str(a) for a in author_list)
                except json.JSONDecodeError:
                    authors_str = authors
            elif isinstance(authors, list):
                authors_str = ','.join(str(a) for a in authors)
            else:
                authors_str = str(authors)

            current_time = datetime.now().astimezone()
            formatted_timestamp = current_time.isoformat(timespec='seconds')

            doc = {
                "id": str(uuid.uuid4()),
                "user": user,
                "framework": framework,
                "accounting_standards": accounting_standards,
                "topics": topics,
                "authors": authors_str,
                "scenario": scenario,
                "query": query,
                "response": response,
                "timestamp": formatted_timestamp,
                "embedding": embedding
            }
            doc_print = {
                "id": str(uuid.uuid4()),
                "user": user,
                "framework": framework,
                "accounting_standards": accounting_standards,
                "topics": topics,
                "authors": authors_str,
                "scenario": scenario,
                "query": query,
                "response": response,
                "timestamp": formatted_timestamp
            }

            print(f"[DEBUG] Document to upload: {doc_print}")
            batch = IndexDocumentsBatch()
            batch.add_merge_or_upload_actions([doc])
            result = self.search_client.index_documents(batch=batch)
            print(f"[DEBUG] Indexing result: {result}")

            return json.dumps({
                "status": "success",
                "message": "Memory exported successfully",
                "document_id": doc["id"]
            })
        except Exception as e:
            print(f"[ERROR] Failed to export memory: {e}")
            return json.dumps({"error": f"Failed to export memory: {str(e)}"})

    @kernel_function(
        description="Retrieve similar conversations from vector store",
        name="retrieve_similar_conversations"
    )
    async def retrieve_similar_conversations(
        self,
        kernel: Kernel,
        arguments: KernelArguments
        ) -> str:
        import json
        try:
            print("[DEBUG] Searching for similar conversations in memory index...")
            print(f"[DEBUG] Arguments: {arguments}")
            user = arguments.get("user", "")
            framework = arguments.get("framework", "")
            accounting_standards = arguments.get("accounting_standards", "")
            topics = arguments.get("topics", "")
            authors = arguments["authors"]
            scenario = arguments["scenario"]
            query = arguments["query"]
            similarity_threshold = arguments.get("similarity_threshold", 1)

            search_text = f"{scenario} {query}"
            vector = await self.get_embedding(search_text)
            print(f"[DEBUG] Embedding for search (first 5 values): {vector[:5]}")

            # Always treat authors as a list
            if isinstance(authors, str):
                try:
                    author_list = json.loads(authors)
                    if not isinstance(author_list, list):
                        author_list = [author_list]
                except json.JSONDecodeError:
                    author_list = [authors]
            else:
                author_list = authors

            # Build filter query in the required order
            filter_clauses = []
            if user:
                filter_clauses.append(f"user eq '{user}'")
                print("[DEBUG] Memory index filter: user applied")
            if framework:
                filter_clauses.append(f"framework eq '{framework}'")
                print("[DEBUG] Memory index filter: framework applied")
            if accounting_standards:
                filter_clauses.append(f"accounting_standards eq '{accounting_standards}'")
                print("[DEBUG] Memory index filter: accounting_standards applied")
            if topics:
                filter_clauses.append(f"topics eq '{topics}'")
                print("[DEBUG] Memory index filter: topics applied")
            if author_list:
                author_filter = " or ".join([f"authors eq '{author}'" for author in author_list])
                filter_clauses.append(f"({author_filter})")
                print("[DEBUG] Memory index filter: authors applied")
            print(f"[DEBUG] Memory index filter order: user, framework, accounting_standards, topics, authors")
            filter_query = " and ".join(filter_clauses)

            results = list(self.search_client.search(
                search_text=search_text,
                filter=filter_query,
                vector_queries=[VectorizedQuery(
                    vector=vector,
                    kind="vector",
                    fields="embedding",
                    k_nearest_neighbors=1
                )],
                select=["authors", "framework", "accounting_standards", "topics", "scenario", "query", "response", "timestamp"],
                top=1,
                query_type="semantic",
                semantic_configuration_name="my-semantic-config",
                query_language="en-us",
                search_fields=["scenario", "query"],
                search_mode="all",
            ))
            print(f"[DEBUG] Azure Search results: {results}")

            if not results:
                print("[DEBUG] No similar conversations found.")
                return json.dumps({"found": False})

            result = results[0]
            # Check for exact match on user, framework, accounting_standards, topics, authors
            exact_match = (
                (user == result.get("user", "")) and
                (framework == result.get("framework", "")) and
                (accounting_standards == result.get("accounting_standards", "")) and
                (topics == result.get("topics", "")) and
                (set(author_list) == set(result.get("authors", "").split(",")))
            )
            if not exact_match:
                print("[DEBUG] Not all filters matched exactly. Generating new response.")
                return json.dumps({"found": False})

            result_text = f"{result['scenario']} {result['query']}"
            result_vector = await self.get_embedding(result_text)
            vector_score = self.calculate_cosine_similarity(vector, result_vector)
            semantic_score = result.get('@search.reranker_score', 0)
            semantic_score /= 4  # Normalize as in your code
            combined_score = (vector_score * 0.7) + (semantic_score * 0.3)
            print(f"[DEBUG] Vector Score: {vector_score}")
            print(f"[DEBUG] Semantic Score: {semantic_score}")
            print(f"[DEBUG] Combined Score: {combined_score}")

            if combined_score > similarity_threshold:
                print("[DEBUG] Similar conversation found above threshold.")
                return json.dumps({
                    "found": True,
                    "conversation": {
                        "authors": result["authors"],
                        "framework": result.get("framework", ""),
                        "accounting_standards": result.get("accounting_standards", ""),
                        "topics": result.get("topics", ""),
                        "scenario": result["scenario"],
                        "query": result["query"],
                        "response": result["response"],
                        "timestamp": result["timestamp"],
                        "similarity_score": combined_score
                    }
                })
            else:
                print("[DEBUG] No similar conversation above threshold.")
                return json.dumps({"found": False})
        except Exception as e:
            print(f"[ERROR] Failed to retrieve similar conversation: {e}")
            return json.dumps({"error": str(e)})

    @kernel_function(
        description="Retrieve all conversations for a given user/session",
        name="retrieve_all_conversations"
    )
    async def retrieve_all_conversations(
        self,
        kernel: Kernel,
        arguments: KernelArguments
    ) -> str:
        """
        Retrieve all conversations for a given user/session.
        Optionally filter by authors.
        """
        try:
            authors = arguments.get("authors", None)
            user = arguments.get("user", None)
            filter_query = None
            if authors:
                if isinstance(authors, str):
                    try:
                        author_list = json.loads(authors)
                    except json.JSONDecodeError:
                        author_list = [authors]
                else:
                    author_list = authors
                filter_query = " or ".join([f"authors eq '{author}'" for author in author_list])
            if user:
                user_filter = f"user eq '{user}'"
                if filter_query:
                    filter_query = f"({user_filter}) and ({filter_query})"
                else:
                    filter_query = user_filter

            results = self.search_client.search(
                search_text="*",
                filter=filter_query,
                select=["authors", "framework", "accounting_standards", "topics", "scenario", "query", "response", "timestamp"],
                top=100,  # Adjust as needed
                query_type="semantic",
                semantic_configuration_name="my-semantic-config",
                query_language="en-us",
                search_fields=["scenario", "query"],
                search_mode="all",
            )
            conversations = []
            for result in results:
                conversations.append({
                    "authors": result.get("authors", []),
                    "framework": result.get("framework", ""),
                    "accounting_standards": result.get("accounting_standards", ""),
                    "topics": result.get("topics", ""),
                    "scenario": result.get("scenario", ""),
                    "query": result.get("query", ""),
                    "question": result.get("query", ""),
                    "response": result.get("response", ""),
                    "timestamp": result.get("timestamp", "")
                })
            return json.dumps({"conversations": conversations})
        except Exception as e:
            return json.dumps({"error": str(e)})
