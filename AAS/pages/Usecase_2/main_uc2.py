# Dependencies and imports
import asyncio
import json
import uuid
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
import re
import os
import pythoncom
from docx2pdf import convert


# Semantic Kernel imports 
from semantic_kernel import Kernel
from semantic_kernel.connectors.ai.open_ai import AzureChatCompletion, AzureTextEmbedding
from semantic_kernel.functions import kernel_function
from semantic_kernel.contents.chat_history import ChatHistory
from semantic_kernel.connectors.ai.function_choice_behavior import FunctionChoiceBehavior
from semantic_kernel.connectors.ai.open_ai.prompt_execution_settings.azure_chat_prompt_execution_settings import (
    AzureChatPromptExecutionSettings,
)

# Azure Search imports
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.models import VectorizedQuery

# Document processing
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration 
@dataclass
class AzureConfig:
    # Azure Search Configuration
    search_endpoint: str = "https://genaipoc-aisearch.search.windows.net"
    search_api_key: str = "WMPSqz89zGLpvssdQgiJVEaDYmvX5TIbJ2QuxQIZCbAzSeDHPy3G"
    
    # Index names for different document types
    guidance_index: str = "guidance-files-index"     # Big 4 guidance + examples
    agreement_index: str = "transactions-docs-index"         # Purchase agreements
    valuation_index: str = "valuation-index"         # Valuation reports  
    financial_index: str = "financial-index"         # Excel balance sheets
    
    # Azure OpenAI Configuration
    openai_api_version: str = "2024-12-01-preview"
    openai_api_key: str = "2rtYSYPHb8h2qMtUSwKtgTvcOn9nj76qwgCVDIUhcmSzDWrkGILLJQQJ99BEACHYHv6XJ3w3AAAAACOGsjHE"
    openai_endpoint: str = "https://ai-premmalothu0362ai608205493981.cognitiveservices.azure.com"
    
    chat_deployment: str = "gpt-4o-mini"
    embedding_deployment: str = "text-embedding-3-large"

# Initialize config
config = AzureConfig()

@dataclass
class BusinessCombinationIssues:
    """Fixed set of 9 business combination issues from the sample"""
    issues: List[str] = None
    
    def __post_init__(self):
        self.issues = [
            "Is the Transaction within the scope of ASC 805?",
            "Which entity is the accounting acquirer?", 
            "What is the acquisition date for accounting purposes?",
            "What is the accounting treatment for Contingent Consideration?",
            "What amounts are included in purchase consideration?",
            "How should the purchase price be allocated, and the assets acquired, and liabilities assumed be recognized on the opening balance sheet?",
            "How should goodwill be recognized and measured?",
            "How will the Company recognize the acquisition related costs incurred as effect of the Acquisition?",
            "How will the Company account for the acquired lease arrangements from the Target?"
        ]

# Test configuration
print("Configuration loaded successfully!")
print(f"Guidance Index: {config.guidance_index}")
print(f"Agreement Index: {config.agreement_index}")

issues = BusinessCombinationIssues()
print(f"Total Issues to Analyze: {len(issues.issues)}")

# Cell 3: Enhanced Document Retrieval Plugin (Works with your actual indexes)
class EnhancedDocumentRetrievalPlugin:
    """Complete document retrieval for all data sources"""
    
    def __init__(self, config: AzureConfig):
        self.config = config
        
        # Initialize search clients for each index
        self.guidance_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.guidance_index,
            credential=AzureKeyCredential(config.search_api_key)
        )
        
        self.agreement_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.agreement_index,
            credential=AzureKeyCredential(config.search_api_key)
        )
        
        # Initialize embedding service
        self.embedding_service = AzureTextEmbedding(
            deployment_name=config.embedding_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
    async def get_embedding(self, text: str) -> List[float]:
        """Get embedding for text"""
        try:
            result = await self.embedding_service.generate_embeddings([text])
            return result[0].data
        except Exception as e:
            print(f"Warning: Embedding generation failed: {e}")
            return [0.0] * 3072
    
    # === GUIDANCE RETRIEVAL FUNCTIONS ===
    @kernel_function(
        description="Retrieve accounting guidance clauses for specific issue",
        name="retrieve_guidance_for_issue"
    )
    async def retrieve_guidance_for_issue(
        self, 
        auditors: str,           # '["KPMG", "Deloitte"]' 
        issue_description: str,   # "ASC 805 scope determination"
        chapter_focus: str = "",  # Optional chapter filter
        top_k: int = 5
    ) -> str:
        """Retrieve technical guidance clauses"""
        
        try:
            # Parse auditors input
            if auditors.startswith('[') and auditors.endswith(']'):
                auditor_list = json.loads(auditors)
            else:
                auditor_list = [auditors]
            
            # Get embedding for issue
            issue_vector = await self.get_embedding(issue_description)
            
            # Build filter for guidance content only
            auditor_filter = " or ".join([f"auditor eq '{auditor}'" for auditor in auditor_list])
            filter_query = f"({auditor_filter}) and content_type eq 'clause'"
            
            if chapter_focus:
                filter_query += f" and chapter_name eq '{chapter_focus}'"
            
            # Create vector query
            vector_query = VectorizedQuery(
                vector=issue_vector,
                kind="vector", 
                fields="embedding",
                k_nearest_neighbors=top_k
            )
            
            # Execute search
            results = self.guidance_client.search(
                search_text=issue_description,
                filter=filter_query,
                vector_queries=[vector_query],
                top=top_k
            )
            
            # Format results
            guidance_docs = []
            for result in results:
                guidance_docs.append({
                    "type": "guidance",
                    "text": result.get("text", ""),
                    "chapter_name": result.get("chapter_name", "Unknown"),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A"),
                    "auditor": result.get("auditor", "Unknown")
                })
            
            print(f"✅ Retrieved {len(guidance_docs)} guidance documents for: {issue_description}")
            return json.dumps(guidance_docs)
            
        except Exception as e:
            error_msg = f"Guidance retrieval error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Retrieve illustrative examples for scenario",
        name="retrieve_examples_for_scenario"
    )
    async def retrieve_examples_for_scenario(
        self, 
        auditors: str,
        scenario_description: str,
        chapter_focus: str = "",
        top_k: int = 3
    ) -> str:
        """Retrieve examples with heading/scenario/assessment structure"""
        
        try:
            # Parse auditors input
            if auditors.startswith('[') and auditors.endswith(']'):
                auditor_list = json.loads(auditors)
            else:
                auditor_list = [auditors]
            
            # Get embedding for scenario
            scenario_vector = await self.get_embedding(scenario_description)
            
            # Build filter for examples only
            auditor_filter = " or ".join([f"auditor eq '{auditor}'" for auditor in auditor_list])
            filter_query = f"({auditor_filter}) and content_type eq 'example'"
            
            if chapter_focus:
                filter_query += f" and chapter_name eq '{chapter_focus}'"
            
            # Create vector query
            vector_query = VectorizedQuery(
                vector=scenario_vector,
                kind="vector",
                fields="embedding", 
                k_nearest_neighbors=top_k
            )
            
            # Execute search
            results = self.guidance_client.search(
                search_text=scenario_description,
                filter=filter_query,
                vector_queries=[vector_query],
                top=top_k
            )
            
            # Format example results
            example_docs = []
            for result in results:
                example_docs.append({
                    "type": "example",
                    "text": result.get("text", ""),
                    "chapter_name": result.get("chapter_name", "Unknown"),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A"),
                    "auditor": result.get("auditor", "Unknown")
                })
            
            print(f"✅ Retrieved {len(example_docs)} examples for: {scenario_description}")
            return json.dumps(example_docs)
            
        except Exception as e:
            error_msg = f"Example retrieval error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})

    # === AGREEMENT RETRIEVAL FUNCTIONS ===
    @kernel_function(
        description="Extract transaction parties from purchase agreement",
        name="extract_transaction_parties"
    )
    async def extract_transaction_parties(self, top_k: int = 10) -> str:
        """Extract buyer, seller, target details from agreement"""
        
        try:
            # Search for parties information
            parties_query = "buyer purchaser seller target company parties between"
            
            results = self.agreement_client.search(
                search_text=parties_query,
                top=top_k
            )
            
            parties_docs = []
            for result in results:
                parties_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(parties_docs)} party documents")
            return json.dumps(parties_docs)
            
        except Exception as e:
            error_msg = f"Parties extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Extract purchase price details from agreement",
        name="extract_purchase_price_details"
    )
    async def extract_purchase_price_details(self, top_k: int = 15) -> str:
        """Extract all purchase price components"""
        
        try:
            # Search for price-related information
            price_query = "purchase price consideration headline ticker contingent additional"
            
            results = self.agreement_client.search(
                search_text=price_query,
                top=top_k
            )
            
            price_docs = []
            for result in results:
                price_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(price_docs)} price documents")
            return json.dumps(price_docs)
            
        except Exception as e:
            error_msg = f"Price extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Extract key contract terms from agreement", 
        name="extract_key_terms"
    )
    async def extract_key_terms(self, top_k: int = 20) -> str:
        """Extract material contract terms"""
        
        try:
            # Search for key terms
            terms_query = "leakage payments closing transaction bonuses employees costs expenses liabilities lease warranties"
            
            results = self.agreement_client.search(
                search_text=terms_query,
                top=top_k
            )
            
            terms_docs = []
            for result in results:
                terms_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(terms_docs)} terms documents")
            return json.dumps(terms_docs)
            
        except Exception as e:
            error_msg = f"Terms extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})

print("✅ Enhanced Document Retrieval Plugin defined!")

# Initialize retrieval plugin
retrieval_plugin = EnhancedDocumentRetrievalPlugin(config)
print("✅ Retrieval plugin initialized successfully!")

# 3-Agent System for Issue Analysis
class MultiAgentIssueAnalysisPlugin:
    """3-agent collaborative system for accounting issue analysis"""
    
    def __init__(self, retrieval_plugin: EnhancedDocumentRetrievalPlugin, config: AzureConfig):
        self.retrieval_plugin = retrieval_plugin
        self.config = config
        
        # Initialize Azure OpenAI service for agents
        self.ai_service = AzureChatCompletion(
            deployment_name=config.chat_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
        # Agent execution settings
        self.execution_settings = AzureChatPromptExecutionSettings(
            max_tokens=3000,
            temperature=0.1
        )
    
    @kernel_function(
        description="Orchestrate 3-agent conversation for accounting issue analysis",
        name="analyze_issue_with_agents"
    )
    async def analyze_issue_with_agents(
        self,
        issue_name: str,
        issue_description: str, 
        transaction_data: str,
        auditors: str = '["KPMG", "Deloitte", "PwC"]'
    ) -> str:
        """Complete 3-agent conversation for issue analysis"""
        
        print(f"Starting 3-agent analysis for: {issue_name}")
        
        try:
            # Retrieve guidance and examples in parallel
            guidance_task = self.retrieval_plugin.retrieve_guidance_for_issue(
                auditors, issue_description, top_k=4
            )
            examples_task = self.retrieval_plugin.retrieve_examples_for_scenario(
                auditors, issue_description, top_k=3
            )
            
            guidance_data, examples_data = await asyncio.gather(guidance_task, examples_task)
            
            # Agent 1: Technical Accountant Analysis
            technical_analysis = await self._technical_accountant_analysis(
                issue_name, issue_description, transaction_data, guidance_data
            )
            
            # Agent 2: Example Researcher Analysis  
            example_analysis = await self._example_researcher_analysis(
                issue_name, issue_description, transaction_data, examples_data
            )
            
            # Agent 3: Synthesizer - Create final conclusion
            final_conclusion = await self._conclusion_synthesizer(
                issue_name, issue_description, transaction_data, 
                technical_analysis, example_analysis, guidance_data, examples_data
            )
            
            # Complete analysis result
            analysis_result = {
                "issue": issue_name,
                "technical_analysis": technical_analysis,
                "example_analysis": example_analysis,
                "final_conclusion": final_conclusion,
                "guidance_retrieved": len(json.loads(guidance_data)) if not "error" in guidance_data else 0,
                "examples_retrieved": len(json.loads(examples_data)) if not "error" in examples_data else 0,
                "timestamp": datetime.now().isoformat()
            }
            
            print(f"3-agent analysis complete for: {issue_name}")
            return json.dumps(analysis_result)
            
        except Exception as e:
            error_msg = f"Multi-agent analysis error: {str(e)}"
            print(f"{error_msg}")
            return json.dumps({"error": error_msg, "issue": issue_name})
    
    async def _technical_accountant_analysis(self, issue_name: str, issue_description: str, 
                                           transaction_data: str, guidance_data: str) -> str:
        """Agent 1: Technical analysis based on authoritative guidance"""
        
        prompt = f"""You are a Senior Technical Accountant analyzing {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Available Guidance**: {guidance_data}

Provide technical analysis covering:
1. **Applicable Standards**: Specific ASC sections that apply
2. **Technical Requirements**: What the guidance requires
3. **Key Criteria**: Critical factors for determination
4. **Application to Facts**: How guidance applies to this transaction
5. **Technical Assessment**: Your technical view based on standards

**Provide detailed technical analysis:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    async def _example_researcher_analysis(self, issue_name: str, issue_description: str,
                                         transaction_data: str, examples_data: str) -> str:
        """Agent 2: Example and precedent analysis"""
        
        prompt = f"""You are an Expert Example Researcher analyzing {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Available Examples**: {examples_data}

Provide example-based analysis covering:
1. **Relevant Examples**: Most applicable precedents from guidance
2. **Fact Pattern Comparison**: How examples compare to our facts
3. **Example Conclusions**: What was determined in similar cases
4. **Precedent Application**: How examples guide our determination
5. **Practical Insights**: Real-world application considerations

**Provide detailed example analysis:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    async def _conclusion_synthesizer(self, issue_name: str, issue_description: str,
                                    transaction_data: str, technical_analysis: str,
                                    example_analysis: str, guidance_data: str, examples_data: str) -> str:
        """Agent 3: Synthesize final memorandum-quality conclusion"""
        
        prompt = f"""You are the Chief Accounting Officer creating the final accounting memorandum section for {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Technical Analysis**: {technical_analysis}
**Example Analysis**: {example_analysis}
**Available Guidance**: {guidance_data}
**Available Examples**: {examples_data}

Create a complete accounting memorandum section following this structure:

**{issue_name}**

**Guidance:**
[Cite relevant ASC sections with specific paragraph numbers]

**Analysis:**
[Integrated analysis combining technical requirements and practical examples]

**Application to Transaction:**
[How the guidance and examples apply to these specific facts]

**Conclusion:**
[Clear, definitive accounting determination with supporting rationale]

**Format as a professional accounting memorandum section with proper ASC citations:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)


# Initialize multi-agent plugin
multi_agent_plugin = MultiAgentIssueAnalysisPlugin(retrieval_plugin, config)


# Response Generation Plugin
class ResponseGenerationPlugin:
    """Generate all sections of the accounting memorandum with dynamic content"""
    
    def __init__(self, retrieval_plugin: EnhancedDocumentRetrievalPlugin, 
                 multi_agent_plugin: MultiAgentIssueAnalysisPlugin, config: AzureConfig):
        self.retrieval_plugin = retrieval_plugin
        self.multi_agent_plugin = multi_agent_plugin
        self.config = config
        
        # Initialize Azure OpenAI service
        self.ai_service = AzureChatCompletion(
            deployment_name=config.chat_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
        self.execution_settings = AzureChatPromptExecutionSettings(
            max_tokens=2000,
            temperature=0.0
        )
    
    # ENTITY EXTRACTION
    @kernel_function(
        description="Extract specific entity names and details from transaction data",
        name="extract_entity_details"
    )
    async def extract_entity_details(self, parties_data: str, price_data: str) -> str:
        """Extract specific entity names and transaction details"""
        
        prompt = f"""Extract the specific entity details from the transaction data:

**Parties Data**: {parties_data}
**Price Data**: {price_data}

Extract and return the specific details in this JSON format:
{{
    "buyer_full_name": "exact buyer company name",
    "buyer_short_name": "short name for buyer",
    "target_full_name": "exact target company name", 
    "target_short_name": "short name for target",
    "target_business_description": "what the target company does",
    "seller_names": "seller entity names",
    "agreement_date": "exact date from documents",
    "agreement_type": "Sale and Purchase Agreement or other",
    "headline_consideration": "exact amount in original currency",
    "ticker_amount": "daily amount if mentioned",
    "target_location": "where target is located",
    "target_operations": "specific business operations details"
}}

**Extract ONLY actual information from the provided data. If information is not available, use 'Not specified' rather than making assumptions:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    # === ENHANCED MEMO HEADER AND SECTIONS ===
    @kernel_function(
        description="Generate memo header and purpose with extracted entity details",
        name="generate_enhanced_memo_header_and_purpose"
    )
    async def generate_enhanced_memo_header_and_purpose(self, entity_details: str) -> str:
        """Generate memo header and purpose using extracted entity details"""
        
        prompt = f"""Generate the memo header and Purpose section using these extracted entity details:

**Entity Details**: {entity_details}

Generate following this exact format using the ACTUAL entity names from the extracted details:

**MEMO HEADER:**
To: Accounting Files
From: Accounting Department  
Subject: Purchase accounting considerations for Acquisition of [USE ACTUAL TARGET NAME]
Date: [Current date: January XX, 2025]
CC: [XXX]

**PURPOSE SECTION:**
The purpose of this memorandum is to evaluate and document the accounting considerations and conclusions reached by the management ("Management") of [USE ACTUAL BUYER FULL NAME] (the "Buyer" or "[USE ACTUAL BUYER SHORT NAME]" or "Company") related to the acquisition (the "Acquisition" or the "Transaction") of [USE ACTUAL TARGET DESCRIPTION] from [USE ACTUAL SELLER NAMES].

The acquisition was consummated pursuant to the [USE ACTUAL AGREEMENT TYPE] dated [USE ACTUAL DATE] entered by and between the Sellers and the Buyer. The analysis is performed in accordance with U.S. GAAP as issued by the Financial Accounting Standards Board ("FASB").

**Use ONLY the actual entity names from the extracted details - NO PLACEHOLDERS like [TARGET NAME]:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate background section with actual business details",
        name="generate_enhanced_background_section"
    )
    async def generate_enhanced_background_section(self, entity_details: str, parties_data: str, price_data: str) -> str:
        """Generate comprehensive background section with actual extracted details"""
        
        prompt = f"""Generate the Background section using ACTUAL details extracted from documents:

**Entity Details**: {entity_details}
**Parties Data**: {parties_data}  
**Price Data**: {price_data}

Create a comprehensive Background section with these subsections using ACTUAL extracted information:

**Target Company Description**: 
- Use actual company name, registered location, business operations
- Include specific business details (e.g., number of stores, business model)
- Mention actual operational focus and market position

**Buyer Company Description**:
- Use actual buyer information including headquarters location  
- Include actual stock exchange, ticker symbol if mentioned in data
- Describe actual business operations and market position

**Acquisition Details**:
- Use actual transaction dates and timeline from the extracted data
- Describe the actual transaction structure (share purchase, asset purchase, etc.)
- Include actual completion timelines and conditions

**Purchase Price Summary**:
- Use actual amounts from the price data with original currency
- Include actual consideration components (headline, ticker, contingent)
- Provide actual transaction structure details

**DO NOT USE any generic placeholders like [insert industry] or [insert location]. Use ONLY the actual extracted information. If specific details are not available, omit those sections rather than using placeholders:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate key terms table from extracted terms",
        name="generate_key_terms_table"
    )
    async def generate_key_terms_table(self, terms_data: str) -> str:
        """Generate key terms section with formatted table (ALREADY WORKING WELL)"""
        
        prompt = f"""Generate the Key Terms section with a formatted table based on retrieved data.

**Terms Data**: {terms_data}

Create a Key Terms section with:
1. **Introduction paragraph** explaining the key terms
2. **Formatted table** with columns: "Term Name and Section" | "Description"
3. **Extract actual terms** like Leakage, Payments at Closing, Transaction Bonuses, etc.

**Format the table with | separators for easy Word conversion:**
| Term Name and Section | Description |
|----------------------|-------------|
| [Term] Section [X] | [Description] |

**Base all terms on the actual retrieved data:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate dynamic literature section based on issues analyzed",
        name="generate_dynamic_literature_section"
    )
    async def generate_dynamic_literature_section(self, all_issue_analyses: str) -> str:
        """Generate accounting literature section based on actual issues analyzed"""
        
        prompt = f"""Generate the Applicable Accounting Literature section based on the issues actually analyzed:

**Issue Analyses**: {all_issue_analyses}

Create a literature section that includes:
1. **Authoritative Literature** - ASC sections actually referenced in the analyses
2. **Non-Authoritative Literature** - Big 4 guidance actually used

**Dynamic structure:**
In performing the analysis of the transaction, Management considered the following authoritative accounting literature:
[List only ASC topics actually referenced in the issue analyses]

In addition to the FASB guidance listed above, Management also considered the following non-authoritative accounting literature in its analysis:
[List only the Big 4 guidance actually retrieved and used]

**Extract the actual ASC sections and guidance sources mentioned in the issue analyses rather than using a generic list:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate documents reviewed section from actual retrieved data",
        name="generate_documents_reviewed_section"
    )
    async def generate_documents_reviewed_section(self, parties_data: str, price_data: str, terms_data: str) -> str:
        """Generate documents reviewed section from actual document sources"""
        
        prompt = f"""Generate the Documents Reviewed section based on actual documents referenced:

**Parties Data**: {parties_data}
**Price Data**: {price_data}
**Terms Data**: {terms_data}

Generate a Documents Reviewed section listing the key documents actually referenced in the data:
1. **Primary Agreement** with actual name and date
2. **Supporting schedules** actually mentioned
3. **Other documents** actually referenced

**Extract actual document names, dates, and references from the provided data rather than using generic examples:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    # === ISSUE ANALYSIS FUNCTIONS (UNCHANGED - ALREADY WORKING) ===
    @kernel_function(
        description="Analyze all 9 business combination issues",
        name="analyze_all_accounting_issues"
    )
    async def analyze_all_accounting_issues(
        self, 
        transaction_data: str,
        parties_data: str, 
        price_data: str,
        terms_data: str
    ) -> str:
        """Analyze all 9 issues using multi-agent system (UNCHANGED)"""
        
        print("Starting analysis of all 9 accounting issues...")
        
        issues = BusinessCombinationIssues().issues
        issue_analyses = []
        
        # Create tasks for parallel execution of issue analyses
        tasks = []
        for i, issue in enumerate(issues, 1):
            task = self.multi_agent_plugin.analyze_issue_with_agents(
                issue_name=f"Issue {i}: {issue}",
                issue_description=issue,
                transaction_data=transaction_data if i <= 3 else 
                               (parties_data if i == 2 else 
                                (price_data if i in [4, 5] else 
                                 (terms_data if i >= 6 else transaction_data)))
            )
            tasks.append(task)
        
        # Execute all issue analyses in parallel
        analyses_results = await asyncio.gather(*tasks)
        
        for result in analyses_results:
            issue_analyses.append(result)
        
        print(f"Completed analysis of all {len(issue_analyses)} issues")
        return json.dumps(issue_analyses)
    
    @kernel_function(
        description="Generate executive summary from all issue analyses",
        name="generate_executive_summary"
    )
    async def generate_executive_summary(self, all_issue_analyses: str) -> str:
        """Generate executive summary with conclusions table (UNCHANGED)"""
        
        prompt = f"""Generate the Executive Summary section based on all issue analyses.

**All Issue Analyses**: {all_issue_analyses}

Create an Executive Summary with:
1. **Introduction paragraph**
2. **Summary table** with columns: "Issue" | "Conclusion"
3. **Extract actual conclusions** from each issue analysis

**Table Format:**
| Issue | Conclusion |
|-------|------------|
| Issue 1: Is the transaction within the scope of ASC 805? | [Actual conclusion] |
| Issue 2: Which entity is the accounting acquirer? | [Actual conclusion] |
[Continue for all 9 issues]

**Extract the actual conclusions from the multi-agent analyses provided:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)


# Initialize response plugin
response_plugin = ResponseGenerationPlugin(retrieval_plugin, multi_agent_plugin, config)

# Cell 6: Complete Document Generation Plugin
# Enhanced Document Generation Plugin with better PDF support
class DocumentGenerationPlugin:
    """Generate final Word documents matching exact sample format"""

    
    def convert_docx_to_pdf_robust(self, docx_path: str, pdf_path: str) -> bool:
        """Robust PDF conversion with multiple fallback methods"""
        try:
            # Method 1: docx2pdf (preferred)
            try:
                import pythoncom
                pythoncom.CoInitialize()
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    print(f"✅ PDF converted successfully using docx2pdf: {pdf_path}")
                    return True
            except Exception as e:
                print(f"❌ docx2pdf failed: {e}")
            
            # # Method 2: Using python-docx and reportlab
            # try:
            #     from docx import Document
            #     from reportlab.lib.pagesizes import letter, A4
            #     from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            #     from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            #     from reportlab.lib.units import inch
            #     from reportlab.lib import colors
                
            #     # Read DOCX content
            #     doc = Document(docx_path)
                
            #     # Create PDF
            #     pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4, 
            #                               rightMargin=72, leftMargin=72,
            #                               topMargin=72, bottomMargin=18)
                
            #     styles = getSampleStyleSheet()
            #     title_style = ParagraphStyle(
            #         'CustomTitle',
            #         parent=styles['Heading1'],
            #         fontSize=16,
            #         spaceAfter=30,
            #     )
                
            #     story = []
                
            #     for para in doc.paragraphs:
            #         if para.text.strip():
            #             # Check if it's a heading (simple heuristic)
            #             if para.style.name.startswith('Heading') or len(para.text) < 100 and para.text.isupper():
            #                 p = Paragraph(para.text, title_style)
            #             else:
            #                 p = Paragraph(para.text, styles['Normal'])
            #             story.append(p)
            #             story.append(Spacer(1, 12))
                
            #     # Handle tables
            #     for table in doc.tables:
            #         table_data = []
            #         for row in table.rows:
            #             row_data = []
            #             for cell in row.cells:
            #                 row_data.append(cell.text)
            #             table_data.append(row_data)
                    
            #         if table_data:
            #             t = Table(table_data)
            #             t.setStyle(TableStyle([
            #                 ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            #                 ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            #                 ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            #                 ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            #                 ('FONTSIZE', (0, 0), (-1, 0), 14),
            #                 ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            #                 ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            #                 ('GRID', (0, 0), (-1, -1), 1, colors.black)
            #             ]))
            #             story.append(t)
            #             story.append(Spacer(1, 12))
                
            #     pdf_doc.build(story)
                
            #     if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            #         print(f"✅ PDF converted successfully using reportlab: {pdf_path}")
            #         return True
                    
            # except Exception as e:
            #     print(f"❌ reportlab conversion failed: {e}")
            
            # # Method 3: Using pypandoc (if available)
            # try:
            #     import pypandoc
            #     pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
            #     if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            #         print(f"✅ PDF converted successfully using pypandoc: {pdf_path}")
            #         return True
            # except Exception as e:
            #     print(f"❌ pypandoc conversion failed: {e}")
            
            # return False
            
        except Exception as e:
            print(f"❌ All PDF conversion methods failed: {e}")
            return False
    
    @kernel_function(
        description="Generate complete accounting memorandum DOCX",
        name="generate_complete_memorandum"
    )
    def generate_complete_memorandum(
        self,
        memo_header_purpose: str,
        background: str,
        key_terms: str,
        literature: str,
        documents_reviewed: str,
        all_issue_analyses: str,
        executive_summary: str,
        filename: str = "accounting_memorandum.docx",
        pdf_filename: str = None
    ) -> dict:
        """Generate complete accounting memorandum DOCX file and PDF, returning file paths"""
        try:
            # Create new document
            doc = Document()
            
            # Parse memo header and purpose
            sections = memo_header_purpose.split("**PURPOSE SECTION:**")
            header_text = sections[0].replace("**MEMO HEADER:**", "").strip()
            purpose_text = sections[1].strip() if len(sections) > 1 else ""
            
            # Add header information
            for line in header_text.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()  # Empty line
            
            # Purpose section
            doc.add_heading("Purpose", level=1)
            doc.add_paragraph(purpose_text)
            
            # Background section
            doc.add_heading("Background", level=1)
            doc.add_paragraph(background)
            
            # Key Terms section with table
            doc.add_heading("Key Terms", level=1)
            self._add_key_terms_with_table(doc, key_terms)
            
            # Literature section
            doc.add_heading("Applicable Accounting Literature and Other Guidance", level=1)
            doc.add_paragraph(literature)
            
            # Documents Reviewed section
            doc.add_heading("Documents Reviewed", level=1)
            doc.add_paragraph(documents_reviewed)
            
            # Accounting Issues section
            doc.add_heading("Accounting Issues", level=1)
            
            # List the 9 issues
            issues = BusinessCombinationIssues().issues
            for i, issue in enumerate(issues, 1):
                doc.add_paragraph(f"Issue {i}: {issue}")
            
            # Add detailed issue analyses
            try:
                issue_analyses = json.loads(all_issue_analyses)
                for analysis_json in issue_analyses:
                    if isinstance(analysis_json, str):
                        analysis = json.loads(analysis_json)
                    else:
                        analysis = analysis_json
                    
                    if "error" not in analysis:
                        # Add issue as heading
                        issue_title = analysis.get('issue', 'Unknown Issue')
                        doc.add_heading(issue_title, level=2)
                        
                        # Add the final conclusion
                        final_conclusion = analysis.get('final_conclusion', 'Analysis pending')
                        doc.add_paragraph(final_conclusion)
                        
            except Exception as e:
                print(f"Warning: Could not parse issue analyses: {e}")
                doc.add_paragraph("Issue analyses will be added after multi-agent processing")
            
            # Executive Summary section
            doc.add_heading("Executive Summary", level=1)
            self._add_executive_summary_with_table(doc, executive_summary)
            
            # Appendices section
            doc.add_heading("Appendices", level=1)
            doc.add_paragraph("Appendix 1: Sale and Purchase Agreement")
            doc.add_paragraph("Appendix 2: Closing Statement and Fund flow statement")
            doc.add_paragraph("Appendix 3: Valuation Report")
            doc.add_paragraph("Appendix 4: Opening Balance Sheet")
            
            # Save DOCX document
            doc.save(filename)
            print(f"✅ DOCX saved: {filename}")
            
            # Generate PDF
            pdf_path = None
            if pdf_filename:
                if self.convert_docx_to_pdf_robust(filename, pdf_filename):
                    pdf_path = pdf_filename
                    print(f"✅ PDF saved: {pdf_filename}")
                else:
                    print(f"❌ PDF conversion failed for: {pdf_filename}")
            
            return {
                "docx_path": filename,
                "pdf_path": pdf_path,
                "message": f"Complete accounting memorandum saved to: {filename}" + (f" and PDF: {pdf_path}" if pdf_path else " (PDF conversion failed)")
            }
            
        except Exception as e:
            error_msg = f"Error generating memorandum: {str(e)}"
            print(error_msg)
            return {
                "error": error_msg
            }
    
    def _add_key_terms_with_table(self, doc, key_terms_text: str):
        """Add key terms section with properly formatted table"""
        
        # Add introductory text
        lines = key_terms_text.split('\n')
        table_started = False
        
        for line in lines:
            line = line.strip()
            if line.startswith('|') and not table_started:
                # Start of table - create table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Term Name and Section"
                header_cells[1].text = "Description"
                
                table_started = True
                continue
            elif line.startswith('|') and table_started:
                # Table content row
                if '|' in line:
                    parts = line.split('|')
                    if len(parts) >= 3:  # Ensure we have content
                        term = parts[1].strip()
                        description = parts[2].strip()
                        
                        if term and description and term != "Term Name and Section":
                            # Add new row
                            row_cells = table.add_row().cells
                            row_cells[0].text = term
                            row_cells[1].text = description
            elif not table_started and line:
                # Regular paragraph before table
                doc.add_paragraph(line)
    
    def _add_executive_summary_with_table(self, doc, executive_summary_text: str):
        """Add executive summary with properly formatted table"""
        
        lines = executive_summary_text.split('\n')
        table_started = False
        
        for line in lines:
            line = line.strip()
            if line.startswith('|') and not table_started:
                # Create summary table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Issue"
                header_cells[1].text = "Conclusion"
                
                table_started = True
                continue
            elif line.startswith('|') and table_started:
                # Table content row
                if '|' in line:
                    parts = line.split('|')
                    if len(parts) >= 3:
                        issue = parts[1].strip()
                        conclusion = parts[2].strip()
                        
                        if issue and conclusion and issue != "Issue":
                            # Add new row
                            row_cells = table.add_row().cells
                            row_cells[0].text = issue
                            row_cells[1].text = conclusion
            elif not table_started and line:
                # Regular paragraph before table
                doc.add_paragraph(line)
# Initialize document plugin
doc_plugin = DocumentGenerationPlugin()

# Cell 7: Enhanced Complete Pipeline Orchestration (FIXED)
class AccountingMemorandumPipeline:
    """Complete pipeline to generate accounting memorandum with dynamic content"""
    
    def __init__(self):
        self.retrieval_plugin = retrieval_plugin
        self.multi_agent_plugin = multi_agent_plugin  
        self.response_plugin = response_plugin
        self.doc_plugin = doc_plugin
    
    async def generate_complete_memorandum(self, filename: str = "accounting_memorandum.docx", pdf_filename: str = None) -> dict:
        """Generate complete accounting memorandum from scratch with dynamic content, return output paths"""
        
        print("Starting complete accounting memorandum generation...")
        
        try:
            # Step 1: Extract all transaction data in parallel
            print("Step 1: Extracting transaction data...")
            
            parties_task = self.retrieval_plugin.extract_transaction_parties()
            price_task = self.retrieval_plugin.extract_purchase_price_details()
            terms_task = self.retrieval_plugin.extract_key_terms()
            
            parties_data, price_data, terms_data = await asyncio.gather(
                parties_task, price_task, terms_task
            )
            
            print(f" Extracted parties, price, and terms data")
            
            # Step 2: Extract entity details for dynamic content generation
            print("Step 2: Extracting entity details...")
            
            entity_details = await self.response_plugin.extract_entity_details(
                parties_data, price_data
            )
            
            print(f" Extracted entity details")
            
            # Step 3: Generate memo sections with dynamic content
            print("Step 3: Generating memo sections with dynamic content...")
            
            header_purpose_task = self.response_plugin.generate_enhanced_memo_header_and_purpose(
                entity_details
            )
            background_task = self.response_plugin.generate_enhanced_background_section(
                entity_details, parties_data, price_data
            )
            key_terms_task = self.response_plugin.generate_key_terms_table(terms_data)
            docs_reviewed_task = self.response_plugin.generate_documents_reviewed_section(
                parties_data, price_data, terms_data
            )
            
            header_purpose, background, key_terms, docs_reviewed = await asyncio.gather(
                header_purpose_task, background_task, key_terms_task, docs_reviewed_task
            )
            
            print(f"   Generated enhanced memo sections")
            
            # Step 4: Analyze all 9 accounting issues (this will take time)
            print("⚖️ Step 4: Analyzing all 9 accounting issues with multi-agent system...")
            print("   (This may take several minutes due to multi-agent conversations)")
            
            all_issue_analyses = await self.response_plugin.analyze_all_accounting_issues(
                price_data, parties_data, price_data, terms_data
            )
            
            print(f"   Completed all issue analyses")
            
            # Step 5: Generate dynamic literature section and executive summary
            print("Step 5: Generating dynamic literature section and executive summary...")
            
            literature_task = self.response_plugin.generate_dynamic_literature_section(
                all_issue_analyses
            )
            executive_summary_task = self.response_plugin.generate_executive_summary(
                all_issue_analyses
            )
            
            literature, executive_summary = await asyncio.gather(
                literature_task, executive_summary_task
            )
            
            print(f"   Generated dynamic literature and executive summary")
            
            # Step 6: Generate final DOCX document
            print("Step 6: Generating final Word document...")
            
            result = self.doc_plugin.generate_complete_memorandum(
                memo_header_purpose=header_purpose,
                background=background,
                key_terms=key_terms,
                literature=literature,
                documents_reviewed=docs_reviewed,
                all_issue_analyses=all_issue_analyses,
                executive_summary=executive_summary,
                filename=filename,
                pdf_filename=pdf_filename
            )
            # result is now a dict with docx_path, pdf_path, message
            return result
        except Exception as e:
            error_msg = f"Pipeline error: {str(e)}"
            print(f"{error_msg}")
            return {"error": error_msg}

# Initialize enhanced complete pipeline
pipeline = AccountingMemorandumPipeline()
