# Cell 1: Complete dependencies and imports
import asyncio
import json
import uuid
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
import re

# Semantic Kernel imports (latest patterns)
from semantic_kernel import Kernel
from semantic_kernel.connectors.ai.open_ai import AzureChatCompletion, AzureTextEmbedding
from semantic_kernel.functions import kernel_function
from semantic_kernel.contents.chat_history import ChatHistory
from semantic_kernel.connectors.ai.function_choice_behavior import FunctionChoiceBehavior
from semantic_kernel.connectors.ai.open_ai.prompt_execution_settings.azure_chat_prompt_execution_settings import (
    AzureChatPromptExecutionSettings,
)

# Azure Search imports
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.models import VectorizedQuery

# Document processing
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("✅ All dependencies imported successfully!")

# Cell 2: Configuration matching your exact setup
@dataclass
class AzureConfig:
    # Azure Search Configuration (from your .env)
    search_endpoint: str = "https://genaipoc-aisearch.search.windows.net"
    search_api_key: str = "WMPSqz89zGLpvssdQgiJVEaDYmvX5TIbJ2QuxQIZCbAzSeDHPy3G"
    
    # Index names for different document types
    guidance_index: str = "guidance-files-index"     # Big 4 guidance + examples
    agreement_index: str = "agreement-index"         # Purchase agreements
    valuation_index: str = "valuation-index"         # Valuation reports  
    financial_index: str = "financial-index"         # Excel balance sheets
    
    # Azure OpenAI Configuration
    openai_api_version: str = "2024-12-01-preview"
    openai_api_key: str = "2rtYSYPHb8h2qMtUSwKtgTvcOn9nj76qwgCVDIUhcmSzDWrkGILLJQQJ99BEACHYHv6XJ3w3AAAAACOGsjHE"
    openai_endpoint: str = "https://ai-premmalothu0362ai608205493981.cognitiveservices.azure.com"
    
    chat_deployment: str = "gpt-4o-mini"
    embedding_deployment: str = "text-embedding-3-large"

# Initialize config
config = AzureConfig()

@dataclass
class BusinessCombinationIssues:
    """Fixed set of 9 business combination issues from the sample"""
    issues: List[str] = None
    
    def __post_init__(self):
        self.issues = [
            "Is the Transaction within the scope of ASC 805?",
            "Which entity is the accounting acquirer?", 
            "What is the acquisition date for accounting purposes?",
            "What is the accounting treatment for Contingent Consideration?",
            "What amounts are included in purchase consideration?",
            "How should the purchase price be allocated, and the assets acquired, and liabilities assumed be recognized on the opening balance sheet?",
            "How should goodwill be recognized and measured?",
            "How will the Company recognize the acquisition related costs incurred as effect of the Acquisition?",
            "How will the Company account for the acquired lease arrangements from the Target?"
        ]

# Test configuration
print("✅ Configuration loaded successfully!")
print(f"�� Guidance Index: {config.guidance_index}")
print(f"�� Agreement Index: {config.agreement_index}")

issues = BusinessCombinationIssues()
print(f"⚖️ Total Issues to Analyze: {len(issues.issues)}")

# Cell 3: Enhanced Document Retrieval Plugin (Works with your actual indexes)
class EnhancedDocumentRetrievalPlugin:
    """Complete document retrieval for all data sources"""
    
    def __init__(self, config: AzureConfig):
        self.config = config
        
        # Initialize search clients for each index
        self.guidance_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.guidance_index,
            credential=AzureKeyCredential(config.search_api_key)
        )
        
        self.agreement_client = SearchClient(
            endpoint=config.search_endpoint,
            index_name=config.agreement_index,
            credential=AzureKeyCredential(config.search_api_key)
        )
        
        # Initialize embedding service
        self.embedding_service = AzureTextEmbedding(
            deployment_name=config.embedding_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
    async def get_embedding(self, text: str) -> List[float]:
        """Get embedding for text"""
        try:
            result = await self.embedding_service.generate_embeddings([text])
            return result[0].data
        except Exception as e:
            print(f"Warning: Embedding generation failed: {e}")
            return [0.0] * 3072
    
    # === GUIDANCE RETRIEVAL FUNCTIONS ===
    @kernel_function(
        description="Retrieve accounting guidance clauses for specific issue",
        name="retrieve_guidance_for_issue"
    )
    async def retrieve_guidance_for_issue(
        self, 
        auditors: str,           # '["KPMG", "Deloitte"]' 
        issue_description: str,   # "ASC 805 scope determination"
        chapter_focus: str = "",  # Optional chapter filter
        top_k: int = 5
    ) -> str:
        """Retrieve technical guidance clauses"""
        
        try:
            # Parse auditors input
            if auditors.startswith('[') and auditors.endswith(']'):
                auditor_list = json.loads(auditors)
            else:
                auditor_list = [auditors]
            
            # Get embedding for issue
            issue_vector = await self.get_embedding(issue_description)
            
            # Build filter for guidance content only
            auditor_filter = " or ".join([f"auditor eq '{auditor}'" for auditor in auditor_list])
            filter_query = f"({auditor_filter}) and content_type eq 'clause'"
            
            if chapter_focus:
                filter_query += f" and chapter_name eq '{chapter_focus}'"
            
            # Create vector query
            vector_query = VectorizedQuery(
                vector=issue_vector,
                kind="vector", 
                fields="embedding",
                k_nearest_neighbors=top_k
            )
            
            # Execute search
            results = self.guidance_client.search(
                search_text=issue_description,
                filter=filter_query,
                vector_queries=[vector_query],
                top=top_k
            )
            
            # Format results
            guidance_docs = []
            for result in results:
                guidance_docs.append({
                    "type": "guidance",
                    "text": result.get("text", ""),
                    "chapter_name": result.get("chapter_name", "Unknown"),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A"),
                    "auditor": result.get("auditor", "Unknown")
                })
            
            print(f"✅ Retrieved {len(guidance_docs)} guidance documents for: {issue_description}")
            return json.dumps(guidance_docs)
            
        except Exception as e:
            error_msg = f"Guidance retrieval error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Retrieve illustrative examples for scenario",
        name="retrieve_examples_for_scenario"
    )
    async def retrieve_examples_for_scenario(
        self, 
        auditors: str,
        scenario_description: str,
        chapter_focus: str = "",
        top_k: int = 3
    ) -> str:
        """Retrieve examples with heading/scenario/assessment structure"""
        
        try:
            # Parse auditors input
            if auditors.startswith('[') and auditors.endswith(']'):
                auditor_list = json.loads(auditors)
            else:
                auditor_list = [auditors]
            
            # Get embedding for scenario
            scenario_vector = await self.get_embedding(scenario_description)
            
            # Build filter for examples only
            auditor_filter = " or ".join([f"auditor eq '{auditor}'" for auditor in auditor_list])
            filter_query = f"({auditor_filter}) and content_type eq 'example'"
            
            if chapter_focus:
                filter_query += f" and chapter_name eq '{chapter_focus}'"
            
            # Create vector query
            vector_query = VectorizedQuery(
                vector=scenario_vector,
                kind="vector",
                fields="embedding", 
                k_nearest_neighbors=top_k
            )
            
            # Execute search
            results = self.guidance_client.search(
                search_text=scenario_description,
                filter=filter_query,
                vector_queries=[vector_query],
                top=top_k
            )
            
            # Format example results
            example_docs = []
            for result in results:
                example_docs.append({
                    "type": "example",
                    "text": result.get("text", ""),
                    "chapter_name": result.get("chapter_name", "Unknown"),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A"),
                    "auditor": result.get("auditor", "Unknown")
                })
            
            print(f"✅ Retrieved {len(example_docs)} examples for: {scenario_description}")
            return json.dumps(example_docs)
            
        except Exception as e:
            error_msg = f"Example retrieval error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})

    # === AGREEMENT RETRIEVAL FUNCTIONS ===
    @kernel_function(
        description="Extract transaction parties from purchase agreement",
        name="extract_transaction_parties"
    )
    async def extract_transaction_parties(self, top_k: int = 10) -> str:
        """Extract buyer, seller, target details from agreement"""
        
        try:
            # Search for parties information
            parties_query = "buyer purchaser seller target company parties between"
            
            results = self.agreement_client.search(
                search_text=parties_query,
                top=top_k
            )
            
            parties_docs = []
            for result in results:
                parties_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(parties_docs)} party documents")
            return json.dumps(parties_docs)
            
        except Exception as e:
            error_msg = f"Parties extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Extract purchase price details from agreement",
        name="extract_purchase_price_details"
    )
    async def extract_purchase_price_details(self, top_k: int = 15) -> str:
        """Extract all purchase price components"""
        
        try:
            # Search for price-related information
            price_query = "purchase price consideration headline ticker contingent additional"
            
            results = self.agreement_client.search(
                search_text=price_query,
                top=top_k
            )
            
            price_docs = []
            for result in results:
                price_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(price_docs)} price documents")
            return json.dumps(price_docs)
            
        except Exception as e:
            error_msg = f"Price extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})
    
    @kernel_function(
        description="Extract key contract terms from agreement", 
        name="extract_key_terms"
    )
    async def extract_key_terms(self, top_k: int = 20) -> str:
        """Extract material contract terms"""
        
        try:
            # Search for key terms
            terms_query = "leakage payments closing transaction bonuses employees costs expenses liabilities lease warranties"
            
            results = self.agreement_client.search(
                search_text=terms_query,
                top=top_k
            )
            
            terms_docs = []
            for result in results:
                terms_docs.append({
                    "text": result.get("text", ""),
                    "source_file": result.get("source_file", "Unknown"),
                    "page": result.get("page", "N/A")
                })
            
            print(f"✅ Retrieved {len(terms_docs)} terms documents")
            return json.dumps(terms_docs)
            
        except Exception as e:
            error_msg = f"Terms extraction error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg})

print("✅ Enhanced Document Retrieval Plugin defined!")

# Initialize retrieval plugin
retrieval_plugin = EnhancedDocumentRetrievalPlugin(config)
print("✅ Retrieval plugin initialized successfully!")

# Cell 4: Streamlined 3-Agent System for Issue Analysis
class MultiAgentIssueAnalysisPlugin:
    """3-agent collaborative system for accounting issue analysis"""
    
    def __init__(self, retrieval_plugin: EnhancedDocumentRetrievalPlugin, config: AzureConfig):
        self.retrieval_plugin = retrieval_plugin
        self.config = config
        
        # Initialize Azure OpenAI service for agents
        self.ai_service = AzureChatCompletion(
            deployment_name=config.chat_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
        # Agent execution settings
        self.execution_settings = AzureChatPromptExecutionSettings(
            max_tokens=3000,
            temperature=0.1
        )
    
    @kernel_function(
        description="Orchestrate 3-agent conversation for accounting issue analysis",
        name="analyze_issue_with_agents"
    )
    async def analyze_issue_with_agents(
        self,
        issue_name: str,
        issue_description: str, 
        transaction_data: str,
        auditors: str = '["KPMG", "Deloitte", "PwC"]'
    ) -> str:
        """Complete 3-agent conversation for issue analysis"""
        
        print(f"�� Starting 3-agent analysis for: {issue_name}")
        
        try:
            # Retrieve guidance and examples in parallel
            guidance_task = self.retrieval_plugin.retrieve_guidance_for_issue(
                auditors, issue_description, top_k=4
            )
            examples_task = self.retrieval_plugin.retrieve_examples_for_scenario(
                auditors, issue_description, top_k=3
            )
            
            guidance_data, examples_data = await asyncio.gather(guidance_task, examples_task)
            
            # Agent 1: Technical Accountant Analysis
            technical_analysis = await self._technical_accountant_analysis(
                issue_name, issue_description, transaction_data, guidance_data
            )
            
            # Agent 2: Example Researcher Analysis  
            example_analysis = await self._example_researcher_analysis(
                issue_name, issue_description, transaction_data, examples_data
            )
            
            # Agent 3: Synthesizer - Create final conclusion
            final_conclusion = await self._conclusion_synthesizer(
                issue_name, issue_description, transaction_data, 
                technical_analysis, example_analysis, guidance_data, examples_data
            )
            
            # Complete analysis result
            analysis_result = {
                "issue": issue_name,
                "technical_analysis": technical_analysis,
                "example_analysis": example_analysis,
                "final_conclusion": final_conclusion,
                "guidance_retrieved": len(json.loads(guidance_data)) if not "error" in guidance_data else 0,
                "examples_retrieved": len(json.loads(examples_data)) if not "error" in examples_data else 0,
                "timestamp": datetime.now().isoformat()
            }
            
            print(f"✅ 3-agent analysis complete for: {issue_name}")
            return json.dumps(analysis_result)
            
        except Exception as e:
            error_msg = f"Multi-agent analysis error: {str(e)}"
            print(f"❌ {error_msg}")
            return json.dumps({"error": error_msg, "issue": issue_name})
    
    async def _technical_accountant_analysis(self, issue_name: str, issue_description: str, 
                                           transaction_data: str, guidance_data: str) -> str:
        """Agent 1: Technical analysis based on authoritative guidance"""
        
        prompt = f"""You are a Senior Technical Accountant analyzing {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Available Guidance**: {guidance_data}

Provide technical analysis covering:
1. **Applicable Standards**: Specific ASC sections that apply
2. **Technical Requirements**: What the guidance requires
3. **Key Criteria**: Critical factors for determination
4. **Application to Facts**: How guidance applies to this transaction
5. **Technical Assessment**: Your technical view based on standards

**Provide detailed technical analysis:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    async def _example_researcher_analysis(self, issue_name: str, issue_description: str,
                                         transaction_data: str, examples_data: str) -> str:
        """Agent 2: Example and precedent analysis"""
        
        prompt = f"""You are an Expert Example Researcher analyzing {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Available Examples**: {examples_data}

Provide example-based analysis covering:
1. **Relevant Examples**: Most applicable precedents from guidance
2. **Fact Pattern Comparison**: How examples compare to our facts
3. **Example Conclusions**: What was determined in similar cases
4. **Precedent Application**: How examples guide our determination
5. **Practical Insights**: Real-world application considerations

**Provide detailed example analysis:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    async def _conclusion_synthesizer(self, issue_name: str, issue_description: str,
                                    transaction_data: str, technical_analysis: str,
                                    example_analysis: str, guidance_data: str, examples_data: str) -> str:
        """Agent 3: Synthesize final memorandum-quality conclusion"""
        
        prompt = f"""You are the Chief Accounting Officer creating the final accounting memorandum section for {issue_name}.

**Issue**: {issue_description}
**Transaction Facts**: {transaction_data}
**Technical Analysis**: {technical_analysis}
**Example Analysis**: {example_analysis}
**Available Guidance**: {guidance_data}
**Available Examples**: {examples_data}

Create a complete accounting memorandum section following this structure:

**{issue_name}**

**Guidance:**
[Cite relevant ASC sections with specific paragraph numbers]

**Analysis:**
[Integrated analysis combining technical requirements and practical examples]

**Application to Transaction:**
[How the guidance and examples apply to these specific facts]

**Conclusion:**
[Clear, definitive accounting determination with supporting rationale]

**Format as a professional accounting memorandum section with proper ASC citations:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)

print("✅ Multi-Agent Issue Analysis Plugin defined!")

# Initialize multi-agent plugin
multi_agent_plugin = MultiAgentIssueAnalysisPlugin(retrieval_plugin, config)
print("✅ Multi-agent plugin initialized!")


# Cell 5: Complete Response Generation Plugin for all memo sections
class ResponseGenerationPlugin:
    """Generate all sections of the accounting memorandum"""
    
    def __init__(self, retrieval_plugin: EnhancedDocumentRetrievalPlugin, 
                 multi_agent_plugin: MultiAgentIssueAnalysisPlugin, config: AzureConfig):
        self.retrieval_plugin = retrieval_plugin
        self.multi_agent_plugin = multi_agent_plugin
        self.config = config
        
        # Initialize Azure OpenAI service
        self.ai_service = AzureChatCompletion(
            deployment_name=config.chat_deployment,
            endpoint=config.openai_endpoint,
            api_key=config.openai_api_key,
            api_version=config.openai_api_version
        )
        
        self.execution_settings = AzureChatPromptExecutionSettings(
            max_tokens=2000,
            temperature=0.0
        )
    
    # === MEMO HEADER AND BASIC SECTIONS ===
    @kernel_function(
        description="Generate memo header and purpose section",
        name="generate_memo_header_and_purpose"
    )
    async def generate_memo_header_and_purpose(self, parties_data: str, transaction_data: str) -> str:
        """Generate memo header and purpose section"""
        
        prompt = f"""Generate the memo header and Purpose section based on the retrieved data.

**Parties Data**: {parties_data}
**Transaction Data**: {transaction_data}

Generate following this exact format:

**MEMO HEADER:**
To: Accounting Files
From: Accounting Department  
Subject: Purchase accounting considerations for Acquisition of [TARGET NAME]
Date: [Current date]
CC: [XXX]

**PURPOSE SECTION:**
The purpose of this memorandum is to evaluate and document the accounting considerations and conclusions reached by the management ("Management") of [BUYER NAME] (the "Buyer" or "[BUYER SHORT]" or "Company") related to the acquisition (the "Acquisition" or the "Transaction") of [TARGET DESCRIPTION] from [SELLERS].

The acquisition was consummated pursuant to the [AGREEMENT TYPE] dated [DATE] entered by and between the Sellers and the Buyer. The analysis is performed in accordance with U.S. GAAP as issued by the Financial Accounting Standards Board ("FASB").

**Extract actual entity names, dates, and transaction details from the provided data:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate background section",
        name="generate_background_section"
    )
    async def generate_background_section(self, parties_data: str, transaction_data: str, price_data: str) -> str:
        """Generate comprehensive background section"""
        
        prompt = f"""Generate the Background section based on retrieved data.

**Parties Data**: {parties_data}
**Transaction Data**: {transaction_data}  
**Price Data**: {price_data}

Generate a comprehensive Background section including:
1. **Target Company Description**: Business, location, operations
2. **Buyer Company Description**: Business, public company details
3. **Acquisition Details**: Transaction overview and timeline
4. **Purchase Price Summary**: High-level consideration structure
5. **Transaction Figures**: Include figure placeholders

**Base the content on actual data retrieved, following the structure from the sample memorandum:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate key terms table",
        name="generate_key_terms_table"
    )
    async def generate_key_terms_table(self, terms_data: str) -> str:
        """Generate key terms section with formatted table"""
        
        prompt = f"""Generate the Key Terms section with a formatted table based on retrieved data.

**Terms Data**: {terms_data}

Create a Key Terms section with:
1. **Introduction paragraph** explaining the key terms
2. **Formatted table** with columns: "Term Name and Section" | "Description"
3. **Extract actual terms** like Leakage, Payments at Closing, Transaction Bonuses, etc.

**Format the table with | separators for easy Word conversion:**
| Term Name and Section | Description |
|----------------------|-------------|
| [Term] Section [X] | [Description] |

**Base all terms on the actual retrieved data:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    @kernel_function(
        description="Generate standard literature section",
        name="generate_literature_section"
    )
    async def generate_literature_section(self) -> str:
        """Generate standard accounting literature section"""
        
        return """Applicable Accounting Literature and Other Guidance

In performing the analysis of the transaction, Management considered the following authoritative accounting literature:
1. ASC Topic No. 805, Business Combinations ("ASC 805")
2. ASC Topic No. 810, Consolidation ("ASC 810")
3. ASC Topic No. 820, Fair Value Measurement ("ASC 820")

In addition to the FASB guidance listed above, Management also considered the following non-authoritative accounting literature in its analysis:
• Deloitte's Roadmap to Accounting for Business Combinations, December 2023 ("Deloitte's Guide")
• KPMG's Handbook on Accounting for income taxes, July 2024 ("KPMG's Handbook")
• Deloitte's Roadmap on Income Taxes, December 2023 ("Deloitte's Guide on Income Taxes")"""
    
    @kernel_function(
        description="Generate documents reviewed section",
        name="generate_documents_reviewed_section"
    )
    async def generate_documents_reviewed_section(self, transaction_data: str) -> str:
        """Generate documents reviewed section"""
        
        prompt = f"""Generate the Documents Reviewed section based on transaction data.

**Transaction Data**: {transaction_data}

Generate a Documents Reviewed section listing the key documents used in the analysis:
1. Sale and Purchase Agreement with actual date
2. Supporting schedules and documents
3. Other relevant documentation

**Extract actual document names and dates from the provided data:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)
    
    # === ISSUE ANALYSIS FUNCTIONS ===
    @kernel_function(
        description="Analyze all 9 business combination issues",
        name="analyze_all_accounting_issues"
    )
    async def analyze_all_accounting_issues(
        self, 
        transaction_data: str,
        parties_data: str, 
        price_data: str,
        terms_data: str
    ) -> str:
        """Analyze all 9 issues using multi-agent system"""
        
        print("�� Starting analysis of all 9 accounting issues...")
        
        issues = BusinessCombinationIssues().issues
        issue_analyses = []
        
        # Create tasks for parallel execution of issue analyses
        tasks = []
        for i, issue in enumerate(issues, 1):
            task = self.multi_agent_plugin.analyze_issue_with_agents(
                issue_name=f"Issue {i}: {issue}",
                issue_description=issue,
                transaction_data=transaction_data if i <= 3 else 
                               (parties_data if i == 2 else 
                                (price_data if i in [4, 5] else 
                                 (terms_data if i >= 6 else transaction_data)))
            )
            tasks.append(task)
        
        # Execute all issue analyses in parallel
        analyses_results = await asyncio.gather(*tasks)
        
        for result in analyses_results:
            issue_analyses.append(result)
        
        print(f"✅ Completed analysis of all {len(issue_analyses)} issues")
        return json.dumps(issue_analyses)
    
    @kernel_function(
        description="Generate executive summary from all issue analyses",
        name="generate_executive_summary"
    )
    async def generate_executive_summary(self, all_issue_analyses: str) -> str:
        """Generate executive summary with conclusions table"""
        
        prompt = f"""Generate the Executive Summary section based on all issue analyses.

**All Issue Analyses**: {all_issue_analyses}

Create an Executive Summary with:
1. **Introduction paragraph**
2. **Summary table** with columns: "Issue" | "Conclusion"
3. **Extract actual conclusions** from each issue analysis

**Table Format:**
| Issue | Conclusion |
|-------|------------|
| Issue 1: Is the transaction within the scope of ASC 805? | [Actual conclusion] |
| Issue 2: Which entity is the accounting acquirer? | [Actual conclusion] |
[Continue for all 9 issues]

**Extract the actual conclusions from the multi-agent analyses provided:**"""

        chat_history = ChatHistory()
        chat_history.add_user_message(prompt)
        
        response = await self.ai_service.get_chat_message_content(
            chat_history=chat_history,
            settings=self.execution_settings
        )
        
        return str(response)

print("✅ Response Generation Plugin defined!")

# Initialize response plugin
response_plugin = ResponseGenerationPlugin(retrieval_plugin, multi_agent_plugin, config)
print("✅ Response generation plugin initialized!")

# Cell 6: Complete Document Generation Plugin
class DocumentGenerationPlugin:
    """Generate final Word documents matching exact sample format"""
    
    @kernel_function(
        description="Generate complete accounting memorandum DOCX",
        name="generate_complete_memorandum"
    )
    def generate_complete_memorandum(
        self,
        memo_header_purpose: str,
        background: str,
        key_terms: str,
        literature: str,
        documents_reviewed: str,
        all_issue_analyses: str,
        executive_summary: str,
        filename: str = "accounting_memorandum.docx"
    ) -> str:
        """Generate complete accounting memorandum DOCX file"""
        
        try:
            # Create new document
            doc = Document()
            
            # Parse memo header and purpose
            sections = memo_header_purpose.split("**PURPOSE SECTION:**")
            header_text = sections[0].replace("**MEMO HEADER:**", "").strip()
            purpose_text = sections[1].strip() if len(sections) > 1 else ""
            
            # Add header information
            for line in header_text.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()  # Empty line
            
            # Purpose section
            doc.add_heading("Purpose", level=1)
            doc.add_paragraph(purpose_text)
            
            # Background section
            doc.add_heading("Background", level=1)
            doc.add_paragraph(background)
            
            # Key Terms section with table
            doc.add_heading("Key Terms", level=1)
            self._add_key_terms_with_table(doc, key_terms)
            
            # Literature section
            doc.add_heading("Applicable Accounting Literature and Other Guidance", level=1)
            doc.add_paragraph(literature)
            
            # Documents Reviewed section
            doc.add_heading("Documents Reviewed", level=1)
            doc.add_paragraph(documents_reviewed)
            
            # Accounting Issues section
            doc.add_heading("Accounting Issues", level=1)
            
            # List the 9 issues
            issues = BusinessCombinationIssues().issues
            for i, issue in enumerate(issues, 1):
                doc.add_paragraph(f"Issue {i}: {issue}")
            
            # Add detailed issue analyses
            try:
                issue_analyses = json.loads(all_issue_analyses)
                for analysis_json in issue_analyses:
                    if isinstance(analysis_json, str):
                        analysis = json.loads(analysis_json)
                    else:
                        analysis = analysis_json
                    
                    if "error" not in analysis:
                        # Add issue as heading
                        issue_title = analysis.get('issue', 'Unknown Issue')
                        doc.add_heading(issue_title, level=2)
                        
                        # Add the final conclusion
                        final_conclusion = analysis.get('final_conclusion', 'Analysis pending')
                        doc.add_paragraph(final_conclusion)
                        
            except Exception as e:
                print(f"Warning: Could not parse issue analyses: {e}")
                doc.add_paragraph("Issue analyses will be added after multi-agent processing")
            
            # Executive Summary section
            doc.add_heading("Executive Summary", level=1)
            self._add_executive_summary_with_table(doc, executive_summary)
            
            # Appendices section
            doc.add_heading("Appendices", level=1)
            doc.add_paragraph("Appendix 1: Sale and Purchase Agreement")
            doc.add_paragraph("Appendix 2: Closing Statement and Fund flow statement")
            doc.add_paragraph("Appendix 3: Valuation Report")
            doc.add_paragraph("Appendix 4: Opening Balance Sheet")
            
            # Save document
            doc.save(filename)
            
            return f"✅ Complete accounting memorandum saved to: {filename}"
            
        except Exception as e:
            return f"❌ Error generating memorandum: {str(e)}"
    
    def _add_key_terms_with_table(self, doc, key_terms_text: str):
        """Add key terms section with properly formatted table"""
        
        # Add introductory text
        lines = key_terms_text.split('\n')
        table_started = False
        
        for line in lines:
            line = line.strip()
            if line.startswith('|') and not table_started:
                # Start of table - create table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Term Name and Section"
                header_cells[1].text = "Description"
                
                table_started = True
                continue
            elif line.startswith('|') and table_started:
                # Table content row
                if '|' in line:
                    parts = line.split('|')
                    if len(parts) >= 3:  # Ensure we have content
                        term = parts[1].strip()
                        description = parts[2].strip()
                        
                        if term and description and term != "Term Name and Section":
                            # Add new row
                            row_cells = table.add_row().cells
                            row_cells[0].text = term
                            row_cells[1].text = description
            elif not table_started and line:
                # Regular paragraph before table
                doc.add_paragraph(line)
    
    def _add_executive_summary_with_table(self, doc, executive_summary_text: str):
        """Add executive summary with properly formatted table"""
        
        lines = executive_summary_text.split('\n')
        table_started = False
        
        for line in lines:
            line = line.strip()
            if line.startswith('|') and not table_started:
                # Create summary table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                header_cells = table.rows[0].cells
                header_cells[0].text = "Issue"
                header_cells[1].text = "Conclusion"
                
                table_started = True
                continue
            elif line.startswith('|') and table_started:
                # Table content row
                if '|' in line:
                    parts = line.split('|')
                    if len(parts) >= 3:
                        issue = parts[1].strip()
                        conclusion = parts[2].strip()
                        
                        if issue and conclusion and issue != "Issue":
                            # Add new row
                            row_cells = table.add_row().cells
                            row_cells[0].text = issue
                            row_cells[1].text = conclusion
            elif not table_started and line:
                # Regular paragraph before table
                doc.add_paragraph(line)

print("✅ Document Generation Plugin defined!")

# Initialize document plugin
doc_plugin = DocumentGenerationPlugin()
print("✅ Document generation plugin initialized!")


# Cell 7: Complete Pipeline Orchestration
class AccountingMemorandumPipeline:
    """Complete pipeline to generate accounting memorandum"""
    
    def __init__(self):
        self.retrieval_plugin = retrieval_plugin
        self.multi_agent_plugin = multi_agent_plugin  
        self.response_plugin = response_plugin
        self.doc_plugin = doc_plugin
    
    async def generate_complete_memorandum(self, filename: str = "accounting_memorandum.docx") -> str:
        """Generate complete accounting memorandum from scratch"""
        
        print("�� Starting complete accounting memorandum generation...")
        
        try:
            # Step 1: Extract all transaction data in parallel
            print("�� Step 1: Extracting transaction data...")
            
            parties_task = self.retrieval_plugin.extract_transaction_parties()
            price_task = self.retrieval_plugin.extract_purchase_price_details()
            terms_task = self.retrieval_plugin.extract_key_terms()
            
            parties_data, price_data, terms_data = await asyncio.gather(
                parties_task, price_task, terms_task
            )
            
            print(f"   ✅ Extracted parties, price, and terms data")
            
            # Step 2: Generate memo sections in parallel
            print("�� Step 2: Generating memo sections...")
            
            header_purpose_task = self.response_plugin.generate_memo_header_and_purpose(
                parties_data, price_data
            )
            background_task = self.response_plugin.generate_background_section(
                parties_data, price_data, terms_data
            )
            key_terms_task = self.response_plugin.generate_key_terms_table(terms_data)
            docs_reviewed_task = self.response_plugin.generate_documents_reviewed_section(price_data)
            
            header_purpose, background, key_terms, docs_reviewed = await asyncio.gather(
                header_purpose_task, background_task, key_terms_task, docs_reviewed_task
            )
            
            # Generate standard literature section
            literature = self.response_plugin.generate_literature_section()
            
            print(f"   ✅ Generated basic memo sections")
            
            # Step 3: Analyze all 9 accounting issues (this will take time)
            print("⚖️ Step 3: Analyzing all 9 accounting issues with multi-agent system...")
            print("   (This may take several minutes due to multi-agent conversations)")
            
            all_issue_analyses = await self.response_plugin.analyze_all_accounting_issues(
                price_data, parties_data, price_data, terms_data
            )
            
            print(f"   ✅ Completed all issue analyses")
            
            # Step 4: Generate executive summary
            print("�� Step 4: Generating executive summary...")
            
            executive_summary = await self.response_plugin.generate_executive_summary(
                all_issue_analyses
            )
            
            print(f"   ✅ Generated executive summary")
            
            # Step 5: Generate final DOCX document
            print("�� Step 5: Generating final Word document...")
            
            result = self.doc_plugin.generate_complete_memorandum(
                memo_header_purpose=header_purpose,
                background=background,
                key_terms=key_terms,
                literature=literature,
                documents_reviewed=docs_reviewed,
                all_issue_analyses=all_issue_analyses,
                executive_summary=executive_summary,
                filename=filename
            )
            
            print(f"�� Pipeline completed successfully!")
            return result
            
        except Exception as e:
            error_msg = f"Pipeline error: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg

# Initialize complete pipeline
print("✅ Complete Accounting Memorandum Pipeline defined!")
pipeline = AccountingMemorandumPipeline()
print("✅ Pipeline initialized and ready to run!")


# Cell 8: Execute the complete pipeline
async def run_complete_pipeline():
    """Run the complete accounting memorandum generation pipeline"""
    
    print("�� Executing complete accounting memorandum generation pipeline...")
    print("=" * 80)
    
    # Generate the complete memorandum
    result = await pipeline.generate_complete_memorandum(
        filename="starbucks_acquisition_memorandum.docx"
    )
    
    print("=" * 80)
    print(f"�� Final Result: {result}")
    
    return result

# Execute the pipeline
import asyncio
final_result = asyncio.run(run_complete_pipeline())

# Cell 6: Test multi-agent system with a sample issue
async def test_multi_agent_system():
    """Test the 4-agent conversation system"""
    
    print("�� Testing 4-agent conversation system...")
    
    # Sample transaction data for testing
    sample_transaction_data = """
    Transaction Overview:
    - Buyer: StarbucksABC Corporation
    - Target: 23.5 Degrees TopcoXYZ Limited  
    - Purchase Price: £133,036,655.42 headline consideration
    - Transaction Date: October 14, 202X4
    - Assets: 110 franchised Starbucks stores across UK
    - Business Model: Drive-thru focused coffee operations
    """
    
    # Test with Issue 1
    test_issue = "Issue 1: ASC 805 Scope Determination"
    test_description = "Determine whether the transaction is within the scope of ASC 805 business combinations"
    
    print(f"\n�� Testing with: {test_issue}")
    print("   This will simulate a full 4-agent conversation...")
    
    try:
        # Run multi-agent analysis
        conversation_result = await multi_agent_plugin.analyze_issue_with_four_agents(
            issue_name=test_issue,
            issue_description=test_description,
            transaction_data=sample_transaction_data,
            auditors='["KPMG"]'
        )
        
        # Parse and display results
        conversation_data = json.loads(conversation_result)
        
        print(f"\n✅ Multi-agent conversation completed!")
        print(f"   Issue: {conversation_data.get('issue')}")
        print(f"   Conversation Complete: {conversation_data.get('conversation_complete')}")
        print(f"   Timestamp: {conversation_data.get('timestamp')}")
        
        # Show structure of conversation
        print(f"\n�� Conversation Structure:")
        for phase, content in conversation_data.items():
            if phase.startswith('phase_'):
                print(f"   ✓ {phase.replace('_', ' ').title()}")
        
        return conversation_data
        
    except Exception as e:
        print(f"❌ Multi-agent test failed: {e}")
        return None

# Run the multi-agent test
import asyncio
# test_result = asyncio.run(test_multi_agent_system())
